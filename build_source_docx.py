#!/usr/bin/env python3
"""
Collect all project source code and export to a single DOCX file for submission.
Run from project root: python build_source_docx.py
"""

import os
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.oxml.ns import qn
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.oxml.ns import qn

# Project root (where this script lives)
ROOT = os.path.dirname(os.path.abspath(__file__))

# Source files to include: (relative path, display name)
SOURCES = [
    ("requirements.txt", "requirements.txt"),
    ("docker/docker-compose.yml", "docker/docker-compose.yml"),
    ("config/filebeat/filebeat-test-simple.yml", "config/filebeat/filebeat-test-simple.yml"),
    ("config/logstash/pipeline.conf", "config/logstash/pipeline.conf"),
    ("scripts/data_extraction.py", "scripts/data_extraction.py"),
    ("scripts/data_preprocessing.py", "scripts/data_preprocessing.py"),
    ("scripts/ml_detector.py", "scripts/ml_detector.py"),
    ("scripts/elasticsearch_writer.py", "scripts/elasticsearch_writer.py"),
    ("scripts/test_elasticsearch_connection.py", "scripts/test_elasticsearch_connection.py"),
    ("scripts/performance_benchmark.py", "scripts/performance_benchmark.py"),
    ("scripts/compare_methods.py", "scripts/compare_methods.py"),
    ("scripts/ml_evaluator.py", "scripts/ml_evaluator.py"),
    ("scripts/false_positive_analyzer.py", "scripts/false_positive_analyzer.py"),
    ("scripts/find_elasticsearch.py", "scripts/find_elasticsearch.py"),
    ("scripts/verify_ml_alerts.py", "scripts/verify_ml_alerts.py"),
]


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"
        if run._element.rPr is not None and run._element.rPr.rFonts is not None:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    return h


def add_code(doc, code, font_size=9):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(code)
    run.font.name = "Courier New"
    if run._element.rPr is not None:
        if run._element.rPr.rFonts is not None:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(font_size)
    return p


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    add_heading(doc, "ELKShield - Source Code", level=0)
    doc.add_paragraph(
        "Hệ thống giám sát an ninh mạng dựa trên ELK Stack và Machine Learning. "
        "Tài liệu này tổng hợp toàn bộ source code chính của dự án."
    )
    doc.add_paragraph()

    for rel_path, display_name in SOURCES:
        full_path = os.path.join(ROOT, rel_path)
        if not os.path.exists(full_path):
            print(f"[SKIP] Not found: {rel_path}")
            continue
        print(f"[ADD] {rel_path}")
        add_heading(doc, display_name, level=1)
        try:
            with open(full_path, "r", encoding="utf-8", errors="replace") as f:
                code = f.read()
        except Exception as e:
            code = f"# Error reading file: {e}"
        # Normalize line endings
        code = code.replace("\r\n", "\n").replace("\r", "\n")
        add_code(doc, code)

    out_path = os.path.join(ROOT, "SOURCE_CODE_ELKShield.docx")
    doc.save(out_path)
    print(f"\n[OK] Saved: {out_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
