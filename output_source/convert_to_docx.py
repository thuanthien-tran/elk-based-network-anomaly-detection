#!/usr/bin/env python3
"""
Script to convert TONG_HOP_CODE.md to DOCX format
"""

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    import re
    import sys
except ImportError:
    print("Installing required packages...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    import re

def add_heading(doc, text, level):
    """Add heading with Vietnamese font support"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    return heading

def add_paragraph(doc, text, style='Normal'):
    """Add paragraph with Vietnamese font support"""
    para = doc.add_paragraph(style=style)
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
    run.font.size = Pt(10)
    return para

def add_code_block(doc, code, language=''):
    """Add code block with monospace font"""
    para = doc.add_paragraph()
    para.style = 'No Spacing'
    run = para.add_run(code)
    run.font.name = 'Courier New'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 128)  # Dark blue
    # Add gray background (requires custom style)
    return para

def remove_emojis(text):
    """Remove emojis from text"""
    import re
    # Remove emoji patterns
    emoji_pattern = re.compile("["
        u"\U0001F600-\U0001F64F"  # emoticons
        u"\U0001F300-\U0001F5FF"  # symbols & pictographs
        u"\U0001F680-\U0001F6FF"  # transport & map symbols
        u"\U0001F1E0-\U0001F1FF"  # flags (iOS)
        u"\U00002702-\U000027B0"
        u"\U000024C2-\U0001F251"
        "]+", flags=re.UNICODE)
    return emoji_pattern.sub(r'', text)

def convert_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to DOCX"""
    print(f"Reading {md_file}...")
    
    try:
        with open(md_file, 'r', encoding='utf-8') as f:
            content = f.read()
    except UnicodeDecodeError:
        with open(md_file, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
    
    # Remove emojis
    content = remove_emojis(content)
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Parse markdown
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block_lines = []
    code_language = ''
    
    while i < len(lines):
        line = lines[i]
        
        # Check for code block start
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                code_text = '\n'.join(code_block_lines)
                add_code_block(doc, code_text, code_language)
                code_block_lines = []
                code_language = ''
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
                code_language = line[3:].strip()
        elif in_code_block:
            code_block_lines.append(line)
        elif line.startswith('# '):
            # H1
            text = remove_emojis(line[2:].strip())
            add_heading(doc, text, 1)
        elif line.startswith('## '):
            # H2
            text = remove_emojis(line[3:].strip())
            add_heading(doc, text, 2)
        elif line.startswith('### '):
            # H3
            text = remove_emojis(line[4:].strip())
            add_heading(doc, text, 3)
        elif line.startswith('**') and line.endswith('**'):
            # Bold text
            para = doc.add_paragraph()
            text = remove_emojis(line[2:-2])
            run = para.add_run(text)
            run.bold = True
        elif line.strip() == '---':
            # Horizontal rule - skip or add paragraph break
            doc.add_paragraph()
        elif line.strip():
            # Regular paragraph
            text = remove_emojis(line)
            add_paragraph(doc, text)
        else:
            # Empty line
            doc.add_paragraph()
        
        i += 1
    
    # Handle remaining code block
    if in_code_block and code_block_lines:
        code_text = '\n'.join(code_block_lines)
        add_code_block(doc, code_text, code_language)
    
    # Save document
    print(f"Saving to {docx_file}...")
    try:
        doc.save(docx_file)
        print(f"Successfully converted to {docx_file}")
    except Exception as e:
        print(f"Error saving: {e}")
        print("Trying alternative method...")
        # Try with different encoding
        import os
        temp_file = docx_file + '.tmp'
        doc.save(temp_file)
        os.rename(temp_file, docx_file)
        print(f"Successfully converted to {docx_file}")

if __name__ == '__main__':
    md_file = 'TONG_HOP_CODE.md'
    docx_file = 'TONG_HOP_CODE.docx'
    
    if len(sys.argv) > 1:
        md_file = sys.argv[1]
    if len(sys.argv) > 2:
        docx_file = sys.argv[2]
    
    try:
        convert_markdown_to_docx(md_file, docx_file)
    except Exception as e:
        print(f"Error: {e}")
        print("\nAlternative: Use online converter or Pandoc")
        print("  - Online: https://www.markdowntoword.com/")
        print("  - Pandoc: pandoc TONG_HOP_CODE.md -o TONG_HOP_CODE.docx")
