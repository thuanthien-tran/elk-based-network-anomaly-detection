# -*- coding: utf-8 -*-
"""Chuyen file Markdown sang Word (.docx). Chay: python md_to_docx.py"""
import os
import re
import sys

try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

def set_font(run, size=12, bold=False):
    run.font.name = "Times New Roman"
    if run._element.rPr is not None and run._element.rPr.rFonts is not None:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold

def add_heading(doc, text, level=1):
    h = doc.add_heading(level=level)
    r = h.add_run(text)
    set_font(r, size=14 if level == 1 else 12, bold=True)
    return h

def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, bold=bold)
    return p

def parse_table(lines):
    """Parse markdown table into list of rows (list of cells)."""
    rows = []
    for line in lines:
        line = line.strip()
        if not line or not line.startswith("|"):
            continue
        cells = [c.strip() for c in line.split("|")[1:-1]]
        if cells and not all(re.match(r"^[-:]+$", c) for c in cells):
            rows.append(cells)
    return rows

def main():
    import argparse

    parser = argparse.ArgumentParser(description="Chuyen file Markdown sang Word (.docx)")
    parser.add_argument(
        "--input",
        dest="md_path",
        default=os.path.join(SCRIPT_DIR, "BAO_CAO_TIEN_DO.md"),
        help="Duong dan file .md",
    )
    parser.add_argument(
        "--output",
        dest="docx_path",
        default=os.path.join(SCRIPT_DIR, "BAO_CAO_TIEN_DO.docx"),
        help="Duong dan file .docx",
    )
    args = parser.parse_args()

    md_path = args.md_path
    docx_path = args.docx_path

    if not os.path.exists(md_path):
        print("Khong tim thay:", md_path)
        return 1

    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    lines = content.split("\n")
    i = 0
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i]
        raw = line

        if line.strip().startswith("|") and "|" in line:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        else:
            if in_table and table_lines:
                rows = parse_table(table_lines)
                if rows:
                    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    t.style = "Table Grid"
                    for ri, row in enumerate(rows):
                        for ci, cell in enumerate(row):
                            if ci < len(t.rows[ri].cells):
                                t.rows[ri].cells[ci].text = cell
                    doc.add_paragraph()
                in_table = False
                table_lines = []

        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), level=1)
        elif line.startswith("## "):
            add_heading(doc, line[3:].strip(), level=2)
        elif line.startswith("### "):
            add_heading(doc, line[4:].strip(), level=3)
        elif line.strip() == "---":
            doc.add_paragraph()
        elif line.strip().startswith("- ") and ("**" in line or not line.strip().startswith("- **")):
            text = line.strip()[2:]
            p = doc.add_paragraph(style="List Bullet")
            # Handle **bold**
            while "**" in text:
                a, _, rest = text.partition("**")
                if a:
                    p.add_run(a)
                b, _, text = rest.partition("**")
                r = p.add_run(b)
                set_font(r, bold=True)
            if text:
                p.add_run(text)
        elif line.strip().startswith("**") and "**" in line:
            p = doc.add_paragraph()
            text = line.strip()
            while "**" in text:
                a, _, rest = text.partition("**")
                if a:
                    r = p.add_run(a)
                    set_font(r)
                b, _, text = rest.partition("**")
                r = p.add_run(b)
                set_font(r, bold=True)
            if text:
                p.add_run(text)
        elif line.strip():
            add_para(doc, line.strip())
        else:
            doc.add_paragraph()

        i += 1

    if in_table and table_lines:
        rows = parse_table(table_lines)
        if rows:
            t = doc.add_table(rows=len(rows), cols=len(rows[0]))
            t.style = "Table Grid"
            for ri, row in enumerate(rows):
                for ci, cell in enumerate(row):
                    if ci < len(t.rows[ri].cells):
                        t.rows[ri].cells[ci].text = cell

    doc.save(docx_path)
    print("Da luu:", docx_path)
    return 0

if __name__ == "__main__":
    sys.exit(main())
