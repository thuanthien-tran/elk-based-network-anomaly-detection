# -*- coding: utf-8 -*-
"""Chuyen file HTML sang Word (.docx). Chay: python html_to_docx.py"""
import os
import re
import sys

try:
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn
except ImportError:
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

def set_font(run, size=12, bold=False):
    run.font.name = "Times New Roman"
    if run._element.rPr is not None and run._element.rPr.rFonts is not None:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold

def add_heading(doc, text, level=1):
    h = doc.add_heading(level=level)
    r = h.add_run(text)
    set_font(r, size=14 if level == 1 else 12, bold=True)
    return h

def add_para(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_font(r)
    return p

def strip_html(html):
    """Remove HTML tags and decode entities."""
    text = re.sub(r"<[^>]+>", " ", html)
    text = re.sub(r"&nbsp;", " ", text)
    text = re.sub(r"&amp;", "&", text)
    text = re.sub(r"&lt;", "<", text)
    text = re.sub(r"&gt;", ">", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text

def main():
    html_path = os.path.join(SCRIPT_DIR, "DU_KIEN_DATASET_VA_PHAT_TRIEN.html")
    docx_path = os.path.join(SCRIPT_DIR, "DU_KIEN_DATASET_VA_PHAT_TRIEN.docx")

    if not os.path.exists(html_path):
        print("Khong tim thay:", html_path)
        return 1

    with open(html_path, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # Extract blocks in document order: (h1|h2|h3|p|li)...content...
    pattern = re.compile(
        r"<(h1|h2|h3|p|li)[^>]*>(.*?)</\1>",
        re.DOTALL | re.IGNORECASE
    )
    blocks = [(m.start(), m.group(1).lower(), strip_html(m.group(2))) for m in pattern.finditer(content)]
    blocks.sort(key=lambda x: x[0])

    for _, tag, text in blocks:
        if not text:
            continue
        if tag == "h1":
            add_heading(doc, text, level=1)
        elif tag == "h2":
            add_heading(doc, text, level=2)
        elif tag == "h3":
            add_heading(doc, text, level=3)
        elif tag == "li":
            add_bullet(doc, text)
        else:
            add_para(doc, text)

    doc.save(docx_path)
    print("Da luu:", docx_path)
    return 0

if __name__ == "__main__":
    sys.exit(main())
