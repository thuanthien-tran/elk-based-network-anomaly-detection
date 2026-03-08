#!/usr/bin/env python3
"""
Tạo file DOCX từ nội dung mô tả chức năng DEMO (cùng nội dung với DEMO_MO_TA_CHUC_NANG.html).
Chạy: python scripts/export_demo_mota_to_docx.py
Cần: pip install python-docx
"""
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
OUT_DOCX = ROOT / "docs" / "DEMO_MO_TA_CHUC_NANG.docx"

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Can pip install python-docx: pip install python-docx")
    sys.exit(1)


def set_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def add_option(doc, num, title, desc):
    p = doc.add_paragraph()
    p.add_run(f"[{num}] ").bold = True
    p.add_run(title + ". ").bold = True
    p.add_run(desc)
    return p


def main():
    doc = Document()
    doc.add_heading("ELKShield Demo – Mô tả chức năng", 0)
    doc.add_paragraph("Tài liệu mô tả từng lựa chọn trong DEMO.bat và cách chọn để demo.")
    doc.add_paragraph()

    set_heading(doc, "1. Menu chính (0–9)", level=1)

    doc.add_heading("Hệ thống (ELK)", level=2)
    add_option(doc, "1", "Reset dữ liệu",
               "Xóa index cũ trong Elasticsearch (test-logs, ml-alerts…) và có thể xóa registry Filebeat. Dùng khi muốn bắt đầu lại từ đầu.")
    add_option(doc, "2", "Chạy Filebeat",
               "Mở cửa sổ Filebeat đọc test.log (Desktop/Documents) và gửi log lên Logstash → Elasticsearch. Bắt buộc trước khi chạy Detection online [7]. Giữ cửa sổ mở, đợi 15–20 giây.")
    add_option(doc, "3", "Kiểm tra index / Mở Kibana",
               "Kiểm tra index trên ES và mở Kibana (http://localhost:5601). Dùng để xem ml-alerts sau khi chạy [7]. Index pattern: ml-alerts-*.")
    doc.add_paragraph()

    doc.add_heading("Tạo log mẫu (cho test detection)", level=2)
    add_option(doc, "4", "Tạo log mẫu – nhập số dòng",
               "Hỏi số dòng normal và attack, ghi vào test.log (Desktop + Documents) với format SSH. Dùng khi muốn tùy chỉnh lượng log.")
    add_option(doc, "5", "Tạo log mẫu mặc định",
               "Ghi sẵn 2 dòng normal + 5 dòng attack vào test.log, không hỏi. Cách nhanh nhất để có log cho demo.")
    doc.add_paragraph()

    doc.add_heading("Train model (offline)", level=2)
    add_option(doc, "6", "Train model (menu con)",
               "Mở menu con chỉ gồm: train unified, chuẩn bị Synthetic, train Russell Mitchell, train từ CSV. Dùng khi cần train model mới hoặc chuẩn bị dữ liệu. Detection online [7] dùng model do [6]-[1] tạo ra.")
    doc.add_paragraph()

    doc.add_heading("Detection (realtime)", level=2)
    add_option(doc, "7", "Detection online",
               "Lấy log từ ES (hoặc fallback đọc test.log) → preprocess → load ssh_attack_model.joblib → predict → ghi ml-alerts. Cần đã chạy [6]-[1] và bật [2] Filebeat. Đây là option chính để demo phát hiện tấn công realtime.")
    doc.add_paragraph()

    doc.add_heading("Demo", level=2)
    add_option(doc, "8", "Demo nhanh",
               "Chạy script Python không qua ELK: đọc dataset có sẵn, predict, in kết quả. Dùng khi không bật Docker/ELK.")
    add_option(doc, "9", "Chạy nhanh",
               "Ghi log → đợi 15s → extract từ ES → preprocess → train (Isolation Forest tại chỗ) → ghi ml-alerts → mở Kibana. Luồng all-in-one, không dùng model unified.")
    add_option(doc, "0", "Thoát", "Thoát chương trình DEMO.bat.")
    doc.add_paragraph()

    set_heading(doc, "2. Menu con [6] Train model", level=1)
    table = doc.add_table(rows=6, cols=3)
    table.style = "Table Grid"
    h = table.rows[0].cells
    h[0].text = "Số"
    h[1].text = "Chức năng"
    h[2].text = "Mô tả"
    for i, (num, name, desc) in enumerate([
        ("1", "Train UNIFIED", "Gộp Synthetic + Russell + Kaggle → train → lưu ssh_attack_model.joblib. Cần đã chạy [2][3][4]. Sau đó dùng [7] Detection online."),
        ("2", "Chuẩn bị Synthetic", "Sinh ~8000 dòng → data/processed/logs.csv. Chỉ tạo dữ liệu, không train. Dùng cho [1] unified."),
        ("3", "Train Russell Mitchell", "auth.log → CSV → preprocess → train → ghi ES. Cần thư mục data\\russellmitchell\\gather."),
        ("4", "Train từ CSV", "A: Russell Mitchell CSV, B: SSH Kaggle, C: File CSV bất kỳ."),
        ("0", "Quay lại", "Về menu chính."),
    ], start=1):
        r = table.rows[i].cells
        r[0].text = num
        r[1].text = name
        r[2].text = desc
    doc.add_paragraph()

    set_heading(doc, "3. Cách chọn để demo", level=1)
    doc.add_heading("Demo Detection online (khuyến nghị)", level=2)
    doc.add_paragraph(
        "[5] Tạo log mặc định → [2] Chạy Filebeat (giữ mở, đợi 15–20s) → [6]→[1] Train unified (nếu chưa có model) → [7] Detection online → [3] Mở Kibana, xem ml-alerts-*"
    )
    doc.add_paragraph("Thứ tự ngắn: 5 → 2 → (6→1 nếu cần) → 7 → 3.")
    doc.add_paragraph()
    doc.add_heading("Demo chỉ ML (không ELK)", level=2)
    doc.add_paragraph("Chọn [8] Demo nhanh.")
    doc.add_paragraph()
    doc.add_heading("Demo all-in-one nhanh", level=2)
    doc.add_paragraph("[5] Tạo log → [2] Filebeat → [9] Chạy nhanh (tự train tại chỗ + ghi ml-alerts + mở Kibana).")
    doc.add_paragraph()
    doc.add_paragraph("ELKShield – SIEM + ML Hybrid. Tài liệu mô tả chức năng DEMO.bat.")

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    print(f"Da ghi: {OUT_DOCX}")
    return 0


if __name__ == "__main__":
    sys.exit(main() or 0)
