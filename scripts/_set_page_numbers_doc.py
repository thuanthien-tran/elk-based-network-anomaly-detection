from copy import deepcopy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

path = Path('docs/DAMH.docx')
doc = Document(path)

idx = None
for i, p in enumerate(doc.paragraphs):
    t = (p.text or '').strip().upper()
    if 'CHƯƠNG I' in t or 'CHUONG I' in t:
        idx = i
        break
if idx is None:
    raise SystemExit('Cannot find CHUONG I heading')

body = doc._element.body
if body.sectPr is None:
    raise SystemExit('No body.sectPr found')

chapter_p = doc.paragraphs[idx]._p

# Build a section-break paragraph before CHUONG I
new_p = OxmlElement('w:p')
pPr = OxmlElement('w:pPr')
sect = deepcopy(body.sectPr)

# Ensure this inserted section ends with next-page break
type_el = sect.find(qn('w:type'))
if type_el is None:
    type_el = OxmlElement('w:type')
    sect.insert(0, type_el)
type_el.set(qn('w:val'), 'nextPage')

pPr.append(sect)
new_p.append(pPr)
chapter_p.addprevious(new_p)

# Start numbering from 1 at section containing CHUONG I (the final section)
final_sect = body.sectPr
pg = final_sect.find(qn('w:pgNumType'))
if pg is None:
    pg = OxmlElement('w:pgNumType')
    final_sect.append(pg)
pg.set(qn('w:start'), '1')

# Try to clear first section footer/header text (if available)
if len(doc.sections) >= 1:
    sec1 = doc.sections[0]
    for para in sec1.footer.paragraphs:
        para.text = ''

if len(doc.sections) >= 2:
    sec2 = doc.sections[1]
    try:
        sec2.footer.is_linked_to_previous = False
    except Exception:
        pass

# Remove temporary script artifacts if rerun-safe and save
out = Path('docs/DAMH_paged.docx')
doc.save(out)
print('Updated:', out)
print('Sections after save:', len(doc.sections))

